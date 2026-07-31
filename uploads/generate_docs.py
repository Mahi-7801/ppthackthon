#!/usr/bin/env python3
"""Generate all hackathon upload documents for SecureSign Innovation Challenge 2026."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Color Palette ──
NAVY = RGBColor(0x00, 0x33, 0x66)
BLUE = RGBColor(0x00, 0x66, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def setup_doc(title_text, subtitle_text=None):
    """Create a new document with standard styling."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # Title
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(28)

    if subtitle_text:
        sub = doc.add_paragraph(subtitle_text)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.runs[0]
        run.font.color.rgb = BLUE
        run.font.size = Pt(14)
        run.font.italic = True

    doc.add_paragraph("")  # spacer
    return doc


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


# ═══════════════════════════════════════════════════════
# DOCUMENT 1: CONCEPT NOTE
# ═══════════════════════════════════════════════════════
def generate_concept_note():
    doc = setup_doc(
        "SecureSign",
        "Universal Type-C DSC Mobile Signing Platform\nSecureSign Innovation Challenge 2026 — Concept Note"
    )

    # Organization info
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Organization", "Vone Digital"),
        ("Applicant", "Mahi Bujji Papa"),
        ("Contact", "pmahi7801@gmail.com | 6301400137"),
        ("Solution", "SecureSign — Universal Type-C DSC Mobile Signing Platform"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.style.font.size = Pt(10)

    doc.add_paragraph("")

    # 1. Problem Statement
    add_heading_styled(doc, "1. Problem Understanding")
    add_body(doc,
        "Digital Signature Certificates (DSC) are legally mandated for signing documents "
        "in India under the IT Act 2000. Currently, using a DSC requires a desktop or laptop "
        "computer with vendor-specific software (ePass, ProxKey, Watchdata utilities). This "
        "creates a critical bottleneck for:"
    )
    add_bullet(doc, "Government field officers who need to sign approval documents on-site")
    add_bullet(doc, "Bank agents processing loan documents at customer locations")
    add_bullet(doc, "Healthcare workers signing prescriptions and certificates remotely")
    add_bullet(doc, "Legal professionals requiring real-time document execution")
    add_bullet(doc, "Enterprise employees working from remote or field locations")

    add_body(doc,
        "The core technical barrier is the absence of a universal, vendor-agnostic SDK that "
        "enables Type-C DSC dongles to work seamlessly with mobile applications on both "
        "Android and iOS, while maintaining full CCA (Controller of Certifying Authorities) "
        "compliance for legally valid digital signatures."
    )

    # 2. Proposed Solution
    add_heading_styled(doc, "2. Proposed Solution")
    add_body(doc,
        "SecureSign is a complete mobile digital signing platform that enables any field "
        "officer to sign documents using a standard Type-C DSC dongle connected to their "
        "Android smartphone via USB OTG. The solution consists of five integrated layers:"
    )
    add_bullet(doc, "Hardware Layer: Type-C DSC dongle (ePass/ProxKey/Watchdata) via USB OTG")
    add_bullet(doc, "Native Layer: Kotlin CCID/PKCS#11 transport for Android USB communication")
    add_bullet(doc, "Bridge Layer: React Native NativeModules exposing sign(), verifyPin(), getCertificate()")
    add_bullet(doc, "Application Layer: React Native screens for the complete signing workflow")
    add_bullet(doc, "Backend Layer: Node.js + Supabase for document management, PAdES assembly, and audit logging")

    # 3. Key Features
    add_heading_styled(doc, "3. Key Features")
    add_bullet(doc, "Plug-and-play: Connect dongle → scan → sign in under 30 seconds")
    add_bullet(doc, "Vendor-agnostic: Works with ePass, ProxKey, Watchdata, and other PKCS#11 tokens")
    add_bullet(doc, "CCA compliant: Private keys never leave hardware, PIN verified on-token, RFC 3161 timestamps")
    add_bullet(doc, "PAdES/CAdES signatures: Legally valid format compliant with Indian IT Act")
    add_bullet(doc, "Full audit trail: Every operation logged with timestamp, user ID, document hash")
    add_bullet(doc, "Direct sharing: Signed PDF shared via email, WhatsApp, or portal upload")

    # 4. CCA Compliance
    add_heading_styled(doc, "4. CCA Compliance")
    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Light Grid Accent 1'
    table2.rows[0].cells[0].text = "CCA Rule"
    table2.rows[0].cells[1].text = "Implementation"
    rules = [
        ("Rule 1: Private Key Security", "Signing happens entirely on the hardware token. Only the signature blob returns to the app — private key never leaves the secure element."),
        ("Rule 2: PIN Verification", "PIN is sent directly to the token via CCID APDU commands. Never stored in app memory, never transmitted to server."),
        ("Rule 3: PAdES with Timestamp", "Signatures assembled in PAdES format with RFC 3161 timestamps from a Time Stamping Authority."),
        ("Rule 4: Retry Limits", "Token enforces maximum 3 PIN attempts. After 3 failures, token locks for security."),
        ("Rule 5: Audit Trail", "Every sign operation logged: user ID, document hash, signature, timestamp, certificate serial, IP address."),
    ]
    for i, (rule, impl) in enumerate(rules):
        table2.rows[i+1].cells[0].text = rule
        table2.rows[i+1].cells[1].text = impl

    # 5. Impact
    add_heading_styled(doc, "5. Expected Impact")
    add_bullet(doc, "Eliminates dependency on desktop computers for DSC signing")
    add_bullet(doc, "Enables real-time document signing from any field location")
    add_bullet(doc, "Reduces document processing time from hours/days to minutes")
    add_bullet(doc, "Supports Digital India initiative for paperless governance")
    add_bullet(doc, "Applicable to government, banking, healthcare, legal, and enterprise sectors")

    # 6. GitHub
    add_heading_styled(doc, "6. Repository & Demo")
    add_body(doc, "GitHub: https://github.com/Mahi-7801/app1234")
    add_body(doc, "Live APK: https://expo.dev/accounts/mahibujjipapas-team/projects/dsc-mobile-signing/builds/8ce0f3a3-39e2-4b36-8c43-3bd61e8b66dc")
    add_body(doc, "Backend API: https://securesign-backend-v2.onrender.com")

    path = os.path.join(OUTPUT_DIR, "SecureSign_Concept_Note.docx")
    doc.save(path)
    print(f"Created: {path}")


# ═══════════════════════════════════════════════════════
# DOCUMENT 2: PRODUCT PRESENTATION
# ═══════════════════════════════════════════════════════
def generate_presentation():
    doc = setup_doc(
        "SecureSign",
        "Product & Solution Presentation — SecureSign Innovation Challenge 2026"
    )

    # Slide 1: Overview
    add_heading_styled(doc, "1. Product Overview")
    add_body(doc,
        "SecureSign is a universal Type-C DSC mobile signing platform that enables "
        "document signing using hardware security tokens directly from Android smartphones. "
        "It replaces the need for desktop computers and vendor-specific software."
    )

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = "Feature"
    table.rows[0].cells[1].text = "Detail"
    features = [
        ("Platform", "Android (React Native + Expo dev-client)"),
        ("Backend", "Node.js + Express + Supabase (PostgreSQL)"),
        ("Token Support", "ePass, ProxKey, Watchdata, TrustKey (PKCS#11)"),
        ("Connection", "USB OTG (Type-C)"),
        ("Signature Format", "PAdES/CAdES with RFC 3161 timestamp"),
        ("Compliance", "CCA Rules 1-5 fully implemented"),
    ]
    for i, (f, d) in enumerate(features):
        table.rows[i+1].cells[0].text = f
        table.rows[i+1].cells[1].text = d

    # Slide 2: User Workflow
    add_heading_styled(doc, "2. User Workflow (8 Steps)")
    steps = [
        "Login — Officer logs into SecureSign app with email/password (Supabase Auth)",
        "Connect Dongle — Plug Type-C DSC dongle via USB OTG → tap 'Scan Again' → app detects vendor name and serial number",
        "Enter PIN — Native secure PIN pad appears → PIN sent directly to hardware token → verified on-chip",
        "Choose Document — Pick approval PDF from phone storage → SHA-256 hash computed on-device",
        "Tap Sign — Hash sent to dongle → token signs with private key (never leaves hardware) → signature returned",
        "Timestamp — Backend submits signature to RFC 3161 Time Stamping Authority → timestamp token received",
        "PAdES Assembly — Backend embeds signature + timestamp into PDF → creates legally valid signed document",
        "Share — Signed PDF downloaded → shared via email/WhatsApp/portal upload directly from field",
    ]
    for i, step in enumerate(steps, 1):
        add_bullet(doc, f"Step {i}: {step}")

    # Slide 3: Architecture
    add_heading_styled(doc, "3. Solution Architecture")
    add_body(doc,
        "Layer 1 — Hardware: Type-C DSC dongle with PKCS#11 interface\n"
        "Layer 2 — Native: Kotlin module (CCID Transport + P11Wrapper) for USB communication\n"
        "Layer 3 — Bridge: React Native NativeModules (listTokens, connectDevice, verifyPin, sign, getCertificate)\n"
        "Layer 4 — App: 7 React Native screens (Splash, Login, Home, PINEntry, DocumentSelect, SignConfirmation, SecureDocument)\n"
        "Layer 5 — Backend: 10 REST API endpoints on Node.js/Express\n"
        "Layer 6 — Database: Supabase PostgreSQL (users, documents, signing_sessions, audit_logs)"
    )

    # Slide 4: Technical Differentiators
    add_heading_styled(doc, "4. Technical Differentiators")
    add_bullet(doc, "Vendor-agnostic: PKCS#11 standard interface works with all major DSC brands")
    add_bullet(doc, "No proprietary software: Replaces vendor-specific desktop utilities")
    add_bullet(doc, "Hardware-level security: Private keys never leave the token's secure element")
    add_bullet(doc, "Real-time signing: Complete workflow under 30 seconds")
    add_bullet(doc, "Immutable audit trail: Every operation logged to Supabase with full traceability")
    add_bullet(doc, "Offline-capable: Document hash computed on-device, signing works without internet")

    # Slide 5: Backend API
    add_heading_styled(doc, "5. Backend API Endpoints (All Verified)")
    api_table = doc.add_table(rows=11, cols=3)
    api_table.style = 'Light Grid Accent 1'
    api_table.rows[0].cells[0].text = "Endpoint"
    api_table.rows[0].cells[1].text = "Method"
    api_table.rows[0].cells[2].text = "Purpose"
    apis = [
        ("/api/signup", "POST", "User registration"),
        ("/api/login", "POST", "User authentication"),
        ("/api/documents", "POST", "Upload document to Supabase"),
        ("/api/documents/:userId", "GET", "List user documents"),
        ("/api/documents/:id/hash", "POST", "Compute document hash"),
        ("/api/signing-sessions", "POST", "Record signing session"),
        ("/api/submit-timestamp", "POST", "RFC 3161 timestamp"),
        ("/api/assemble-signature", "POST", "PAdES PDF assembly"),
        ("/api/audit-logs", "POST", "Log audit event"),
        ("/api/verify-signature", "POST", "Verify signature"),
    ]
    for i, (ep, method, purpose) in enumerate(apis):
        api_table.rows[i+1].cells[0].text = ep
        api_table.rows[i+1].cells[1].text = method
        api_table.rows[i+1].cells[2].text = purpose

    # Slide 6: Demo
    add_heading_styled(doc, "6. Live Demo Links")
    add_body(doc, "GitHub Repository: https://github.com/Mahi-7801/app1234")
    add_body(doc, "Android APK: https://expo.dev/accounts/mahibujjipapas-team/projects/dsc-mobile-signing/builds/8ce0f3a3-39e2-4b36-8c43-3bd61e8b66dc")
    add_body(doc, "Backend API: https://securesign-backend-v2.onrender.com")
    add_body(doc, "Supabase Dashboard: https://vpgvqzpdvreylcujmmvu.supabase.co")

    path = os.path.join(OUTPUT_DIR, "SecureSign_Product_Presentation.docx")
    doc.save(path)
    print(f"Created: {path}")


# ═══════════════════════════════════════════════════════
# DOCUMENT 3: TECHNICAL ARCHITECTURE
# ═══════════════════════════════════════════════════════
def generate_architecture():
    doc = setup_doc(
        "SecureSign",
        "Technical Architecture Document — SecureSign Innovation Challenge 2026"
    )

    # 1. System Overview
    add_heading_styled(doc, "1. System Overview")
    add_body(doc,
        "SecureSign is a multi-layered system that bridges hardware security tokens with "
        "mobile applications. The architecture follows a strict separation of concerns: "
        "hardware communication is handled by native Kotlin modules, business logic by "
        "React Native, and data persistence by a Node.js backend with Supabase."
    )

    # 2. Architecture Layers
    add_heading_styled(doc, "2. Architecture Layers")

    add_heading_styled(doc, "2.1 Hardware Layer", level=2)
    add_body(doc,
        "Type-C DSC dongles conforming to CCID (Chip Card Interface Device) class 0x0B. "
        "Each token contains a PKCS#11 interface exposing: C_Initialize, C_OpenSession, "
        "C_Login (PIN verify), C_Sign (sign hash), C_GetAttributeValue (read certificate). "
        "Supported vendors: ePass (Feitian), ProxKey (ProxKey), Watchdata, TrustKey."
    )

    add_heading_styled(doc, "2.2 Native Layer (Kotlin)", level=2)
    add_body(doc,
        "Five Kotlin modules handle all hardware communication:"
    )
    add_bullet(doc, "DSCUsbManager.kt — Scans USB devices, requests permission, opens connections")
    add_bullet(doc, "CcidTransport.kt — Sends/receives CCID APDU commands over USB bulk endpoints")
    add_bullet(doc, "P11Wrapper.kt — PKCS#11 wrapper: initialize, verifyPin, sign, getCertificate")
    add_bullet(doc, "DSCSigningModule.kt — React Native bridge: exposes listTokens, connectDevice, verifyPin, sign, getCertificate, disconnect")
    add_bullet(doc, "DSCSigningPackage.kt — Registers the native module with React Native")

    add_heading_styled(doc, "2.3 Bridge Layer (React Native)", level=2)
    add_body(doc,
        "DSCService.ts provides a TypeScript interface to the native modules. "
        "All native calls are wrapped in async functions with error handling. "
        "Event listeners (onDeviceConnected, onDeviceDisconnected) use NativeEventEmitter."
    )

    add_heading_styled(doc, "2.4 Application Layer (React Native)", level=2)
    app_table = doc.add_table(rows=8, cols=3)
    app_table.style = 'Light Grid Accent 1'
    app_table.rows[0].cells[0].text = "Screen"
    app_table.rows[0].cells[1].text = "Purpose"
    app_table.rows[0].cells[2].text = "Key Functions"
    screens = [
        ("SplashScreen", "App loading", "Animation + navigation"),
        ("LoginScreen", "Authentication", "BackendService.login()"),
        ("HomeScreen", "Dongle management", "DSCService.listTokens(), connectDevice()"),
        ("PINEntryScreen", "Token PIN verification", "DSCService.verifyPin()"),
        ("DocumentSelectScreen", "Document picker", "DocumentPicker, hash computation"),
        ("SignConfirmationScreen", "Signing + audit", "DSCService.sign(), BackendService.submitTimestamp()"),
        ("SecureDocumentScreen", "Download signed PDF", "FileSystem.downloadAsync(), Sharing.shareAsync()"),
    ]
    for i, (s, p, f) in enumerate(screens):
        app_table.rows[i+1].cells[0].text = s
        app_table.rows[i+1].cells[1].text = p
        app_table.rows[i+1].cells[2].text = f

    add_heading_styled(doc, "2.5 Backend Layer (Node.js + Supabase)", level=2)
    add_body(doc,
        "Express.js REST API with 10 endpoints. Supabase provides PostgreSQL database, "
        "JWT authentication, and row-level security. All endpoints require Bearer token "
        "authentication via the requireAuth middleware."
    )

    add_heading_styled(doc, "2.6 Database Schema (Supabase PostgreSQL)", level=2)
    db_table = doc.add_table(rows=5, cols=2)
    db_table.style = 'Light Grid Accent 1'
    db_table.rows[0].cells[0].text = "Table"
    db_table.rows[0].cells[1].text = "Key Columns"
    tables = [
        ("users", "id, email, full_name, created_at"),
        ("documents", "id, user_id, document_name, document_hash, storage_path"),
        ("signing_sessions", "id, user_id, document_id, certificate_serial_number, signed_hash, signature_blob, timestamp_token"),
        ("audit_logs", "id, user_id, event_type, event_details, ip_address, timestamp"),
    ]
    for i, (t, c) in enumerate(tables):
        db_table.rows[i+1].cells[0].text = t
        db_table.rows[i+1].cells[1].text = c

    # 3. Security Architecture
    add_heading_styled(doc, "3. Security Architecture")
    add_bullet(doc, "Private Key Isolation: Signing occurs entirely on the hardware token's secure element. The app only receives the signature blob — the private key never crosses the USB boundary.")
    add_bullet(doc, "PIN Security: PIN bytes are sent directly to the token via CCID APDU, zeroed from memory after use (pinBytes.fill(0)), never logged or transmitted to backend.")
    add_bullet(doc, "Transport Security: All backend communication over HTTPS with Bearer JWT tokens. JWT issued by Supabase Auth with 1-hour expiry.")
    add_bullet(doc, "Session Management: 5-minute session timeout. PIN re-verification required after timeout or app backgrounding. Session invalidated on logout.")
    add_bullet(doc, "Data Integrity: Document SHA-256 hash computed on-device before signing. Hash stored immutably in Supabase. Any document modification breaks hash verification.")
    add_bullet(doc, "Audit Trail: Every operation (connect, PIN verify, sign, timestamp, download) logged with timestamp, user ID, document ID, and IP address.")

    # 4. Signing Flow
    add_heading_styled(doc, "4. Digital Signature Flow")
    add_body(doc,
        "1. User picks document → SHA-256 hash computed on-device\n"
        "2. Hash sent to DSC token via DSCService.sign(hash, 'SHA256WithRSA')\n"
        "3. Token signs hash with private key (RSA, PKCS#1 v1.5)\n"
        "4. Signature blob returned as hex string\n"
        "5. Backend submits signature + hash to RFC 3161 TSA\n"
        "6. TSA returns timestamp token (hash of signature + time)\n"
        "7. Backend assembles PAdES signature container\n"
        "8. Signed PDF generated with embedded signature + timestamp\n"
        "9. Audit log recorded: user, document, hash, signature, timestamp, certificate serial"
    )

    # 5. Deployment
    add_heading_styled(doc, "5. Deployment Architecture")
    add_bullet(doc, "Mobile App: Built via EAS Build (Expo Application Services), distributed as APK")
    add_bullet(doc, "Backend: Deployed on Render (Node.js free tier), auto-deploys from GitHub main branch")
    add_bullet(doc, "Database: Supabase hosted PostgreSQL (vpgvqzpdvreylcujmmvu.supabase.co)")
    add_bullet(doc, "Auth: Supabase Auth with email/password + JWT tokens")

    # 6. Links
    add_heading_styled(doc, "6. Resources")
    add_body(doc, "GitHub: https://github.com/Mahi-7801/app1234")
    add_body(doc, "APK: https://expo.dev/accounts/mahibujjipapas-team/projects/dsc-mobile-signing/builds/8ce0f3a3-39e2-4b36-8c43-3bd61e8b66dc")
    add_body(doc, "Backend: https://securesign-backend-v2.onrender.com")

    path = os.path.join(OUTPUT_DIR, "SecureSign_Technical_Architecture.docx")
    doc.save(path)
    print(f"Created: {path}")


# ═══════════════════════════════════════════════════════
# RUN ALL
# ═══════════════════════════════════════════════════════
if __name__ == "__main__":
    generate_concept_note()
    generate_presentation()
    generate_architecture()
    print("\nAll 3 documents generated in:", OUTPUT_DIR)
